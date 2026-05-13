from __future__ import annotations

import importlib.util
from pathlib import Path

from .core.renderer import create_minimal_docx
from .core.settings import PACKAGE_ROOT
from .core.simple_xlsx import write_xlsx

SAMPLE_WORKBOOK = PACKAGE_ROOT / "sample_data" / "Occam_Data.xlsx"
SAMPLE_DOCX = PACKAGE_ROOT / "sample_templates" / "Documents" / "Individual Tax Engagement Letter.docx"

SETTINGS_ROWS = [
    {"Setting": "Firm Name", "Value": "Occam Advisors"},
    {"Setting": "Default Tax Year", "Value": "2025"},
    {"Setting": "Office Phone", "Value": "(312) 555-0198"},
    {"Setting": "Reply Email", "Value": "ops@occamadvisors.example"},
]

CLIENT_ROWS = [
    {"Client ID": "C-1001", "Client Name": "Arbor & Finch LLC", "Client Email": "alex@arborfinch.example", "Contact First Name": "Alex", "Entity Type": "LLC", "Tax Year": "2025", "Fee Amount": "1850", "Partner": "Dana Walsh", "Manager": "Mia Chen", "Billing Terms": "Due on receipt", "Payment Link": "https://pay.example/arbor"},
    {"Client ID": "C-1002", "Client Name": "Riverbend Family Trust", "Client Email": "trustees@riverbend.example", "Contact First Name": "Morgan", "Entity Type": "Trust", "Tax Year": "2025", "Fee Amount": "2400", "Partner": "Dana Walsh", "Manager": "Owen Patel", "Billing Terms": "Net 15", "Payment Link": "https://pay.example/riverbend"},
    {"Client ID": "C-1003", "Client Name": "Northstar Design Co.", "Client Email": "billing@northstardesign.example", "Contact First Name": "Priya", "Entity Type": "S-Corp", "Tax Year": "2025", "Fee Amount": "3200", "Partner": "Leo Grant", "Manager": "Mia Chen", "Billing Terms": "Net 10", "Payment Link": "https://pay.example/northstar"},
    {"Client ID": "C-1004", "Client Name": "Hale Household", "Client Email": "hale@example.net", "Contact First Name": "Jordan", "Entity Type": "Individual", "Tax Year": "2025", "Fee Amount": "750", "Partner": "Leo Grant", "Manager": "Owen Patel", "Billing Terms": "Due on receipt", "Payment Link": "https://pay.example/hale"},
    {"Client ID": "C-1005", "Client Name": "Willow Creek Ventures", "Client Email": "", "Contact First Name": "Sam", "Entity Type": "Partnership", "Tax Year": "2025", "Fee Amount": "0", "Partner": "Dana Walsh", "Manager": "Mia Chen", "Billing Terms": "Net 15", "Payment Link": ""},
]

INVOICE_ROWS = [
    {"Invoice Number": "INV-24017", "Client ID": "C-1001", "Invoice Amount": "1850", "Invoice Date": "2026-05-01", "Due Date": "2026-05-20", "Service Description": "2025 tax engagement setup", "Payment Link": "https://pay.example/arbor/inv-24017"},
    {"Invoice Number": "INV-24018", "Client ID": "C-1002", "Invoice Amount": "2400", "Invoice Date": "2026-05-03", "Due Date": "2026-05-17", "Service Description": "trust tax compliance retainer", "Payment Link": "https://pay.example/riverbend/inv-24018"},
    {"Invoice Number": "INV-24019", "Client ID": "C-1003", "Invoice Amount": "3200", "Invoice Date": "2026-04-22", "Due Date": "2026-05-02", "Service Description": "business tax advisory and return preparation", "Payment Link": "https://pay.example/northstar/inv-24019"},
]

MISSING_ITEM_ROWS = [
    {"Client ID": "C-1001", "Missing Item": "Signed engagement letter", "Status": "Requested", "Notes": "Sent initial request"},
    {"Client ID": "C-1001", "Missing Item": "December bank statement", "Status": "Requested", "Notes": "Needed for reconciliation"},
    {"Client ID": "C-1002", "Missing Item": "Trust distribution schedule", "Status": "Requested", "Notes": ""},
    {"Client ID": "C-1002", "Missing Item": "Brokerage 1099 package", "Status": "Received", "Notes": "Uploaded to portal"},
    {"Client ID": "C-1003", "Missing Item": "Payroll tax filings Q4", "Status": "Requested", "Notes": "Follow up next week"},
    {"Client ID": "C-1003", "Missing Item": "Shareholder basis schedule", "Status": "Requested", "Notes": ""},
    {"Client ID": "C-1004", "Missing Item": "Charitable contribution receipts", "Status": "Requested", "Notes": ""},
    {"Client ID": "C-1005", "Missing Item": "Partner capital rollforward", "Status": "Requested", "Notes": "Validation demo client"},
]

SAMPLE_SHEETS = {
    "Settings": SETTINGS_ROWS,
    "Clients": CLIENT_ROWS,
    "Invoices": INVOICE_ROWS,
    "Missing Items": MISSING_ITEM_ROWS,
}

ENGAGEMENT_PARAGRAPHS = [
    "Occam Advisors",
    "Individual Tax Engagement Letter",
    "Client: {{Client Name}}",
    "Tax Year: {{Tax Year}}",
    "Dear {{Contact First Name}},",
    "Thank you for selecting {{Firm Name}} to assist with your {{Tax Year}} tax matters.",
    "Our fee for this engagement is ${{Fee Amount}}. Billing terms: {{Billing Terms}}.",
    "Engagement partner: {{Partner}}. Manager: {{Manager}}.",
    "Please use this secure payment link if applicable: {{Payment Link}}.",
]


def _has_module(module_name: str) -> bool:
    return importlib.util.find_spec(module_name) is not None


def create_sample_workbook(path: str | Path = SAMPLE_WORKBOOK, force: bool = False) -> Path:
    path = Path(path)
    if path.exists() and not force:
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    if _has_module("openpyxl"):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        from openpyxl.utils import get_column_letter

        workbook = Workbook()
        workbook.remove(workbook.active)
        header_fill = PatternFill("solid", fgColor="E5E7EB")
        title_font = Font(bold=True, color="0B1F3A")
        for sheet_name, rows in SAMPLE_SHEETS.items():
            ws = workbook.create_sheet(sheet_name)
            headers = list(rows[0].keys())
            ws.append(headers)
            for cell in ws[1]:
                cell.font = title_font
                cell.fill = header_fill
            for row in rows:
                ws.append([row.get(header, "") for header in headers])
            ws.freeze_panes = "A2"
            for index, header in enumerate(headers, 1):
                width = max(len(header) + 4, *(len(str(row.get(header, ""))) + 2 for row in rows))
                ws.column_dimensions[get_column_letter(index)].width = min(width, 42)
        workbook.save(path)
    else:
        write_xlsx(path, SAMPLE_SHEETS)
    return path


def create_sample_docx(path: str | Path = SAMPLE_DOCX, force: bool = False) -> Path:
    path = Path(path)
    if path.exists() and not force:
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    if _has_module("docx"):
        from docx import Document
        from docx.shared import Pt, RGBColor

        document = Document()
        title = document.add_paragraph()
        run = title.add_run("Occam Advisors")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(11, 31, 58)
        subtitle = document.add_paragraph()
        subtitle_run = subtitle.add_run("Individual Tax Engagement Letter")
        subtitle_run.bold = True
        for paragraph in ENGAGEMENT_PARAGRAPHS[2:]:
            document.add_paragraph(paragraph)
        document.save(path)
    else:
        create_minimal_docx(path, ENGAGEMENT_PARAGRAPHS)
    return path


def ensure_sample_assets(force: bool = False) -> dict[str, str]:
    workbook = create_sample_workbook(force=force)
    docx = create_sample_docx(force=force)
    return {"workbook": str(workbook), "docx_template": str(docx)}


if __name__ == "__main__":
    assets = ensure_sample_assets(force=True)
    print(f"Generated sample workbook: {assets['workbook']}")
    print(f"Generated sample docx template: {assets['docx_template']}")
