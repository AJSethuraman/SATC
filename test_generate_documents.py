#!/usr/bin/env python3
"""Tests for the local document template generator.

Pure rendering and a temp-folder end-to-end run. No real client data. These do
not require the heavy runtime dependencies (no PyMuPDF/pandas needed).
"""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

import generate_documents as gd

try:
    import docx  # python-docx
    import docxtpl  # noqa: F401

    _HAVE_DOCX = True
except Exception:  # pragma: no cover - depends on environment
    _HAVE_DOCX = False


class RenderTemplateTests(unittest.TestCase):
    def test_simple_field_substitution(self) -> None:
        out = gd.render_template("Hello {{client_name}} ({{tax_year}})",
                                 {"client_name": "Jordan Sample", "tax_year": "2024"})
        self.assertEqual(out, "Hello Jordan Sample (2024)")

    def test_missing_field_renders_blank(self) -> None:
        self.assertEqual(gd.render_template("A{{nope}}B", {}), "AB")

    def test_values_are_html_escaped(self) -> None:
        out = gd.render_template("{{x}}", {"x": "<b>&"})
        self.assertEqual(out, "&lt;b&gt;&amp;")

    def test_single_braces_left_untouched(self) -> None:
        css = "body { color: red; } {{client_name}}"
        self.assertEqual(gd.render_template(css, {"client_name": "X"}), "body { color: red; } X")

    def test_repeating_section(self) -> None:
        template = "<ul>{{#items}}<li>{{description}}: {{amount}}</li>{{/items}}</ul>"
        context = {"items": [
            {"description": "Federal", "amount": "300.00"},
            {"description": "State", "amount": "100.00"},
        ]}
        self.assertEqual(
            gd.render_template(template, context),
            "<ul><li>Federal: 300.00</li><li>State: 100.00</li></ul>",
        )

    def test_empty_section_renders_nothing(self) -> None:
        self.assertEqual(gd.render_template("X{{#items}}<li>{{a}}</li>{{/items}}Y", {"items": []}), "XY")

    def test_missing_fields_detection_ignores_section_internals(self) -> None:
        template = "{{firm_name}} {{#items}}{{description}}{{/items}} {{client_name}}"
        missing = gd.missing_fields(template, {"firm_name": "Firm"})
        self.assertEqual(missing, ["client_name"])  # description is section-internal

    def test_template_selection_semantics(self) -> None:
        directory = gd.REPO_TEMPLATE_DIR
        self.assertEqual(set(gd.available_templates(directory, None)), set(gd.SHIPPED_TEMPLATE_KEYS))
        self.assertEqual(gd.available_templates(directory, []), {})  # explicit empty = none
        self.assertEqual(set(gd.available_templates(directory, ["invoice"])), {"invoice"})

    def test_discovers_dropped_template_files(self) -> None:
        with tempfile.TemporaryDirectory() as d:
            templates = Path(d) / "document_templates"
            templates.mkdir()
            (templates / "thank_you_note.html").write_text("<p>{{client_name}}</p>", encoding="utf-8")
            discovered = gd.available_templates(templates)
            self.assertIn("thank_you_note", discovered)  # no code change needed

    @unittest.skipUnless(_HAVE_DOCX, "docxtpl not installed")
    def test_renders_word_template(self) -> None:
        import docx  # python-docx

        with tempfile.TemporaryDirectory() as d:
            folder = Path(d)
            templates = folder / "document_templates"
            templates.mkdir()
            # Author a tiny Word template with a placeholder.
            source = docx.Document()
            source.add_paragraph("Dear {{ client_name }}, thank you.")
            source.save(templates / "thank_you.docx")
            (folder / "clients.json").write_text(
                json.dumps([{"client_name": "Pat Example"}]), encoding="utf-8"
            )

            result = gd.run_generation(folder, templates=["thank_you"])
            self.assertEqual(result["document_count"], 1)
            out = Path(result["documents"][0])
            self.assertEqual(out.suffix, ".docx")
            rendered = docx.Document(out)
            text = "\n".join(p.text for p in rendered.paragraphs)
            self.assertIn("Pat Example", text)
            self.assertNotIn("{{", text)

    def test_results_letter_renders_table_and_payment(self) -> None:
        directory = gd.REPO_TEMPLATE_DIR
        template = (directory / "tax_results_letter.html").read_text(encoding="utf-8")
        context = gd.augment_context({
            "client_name": "Sam Sample", "tax_year": "2025", "total": "450.00",
            "pay_link": "https://pay.example.com/sam",
            "returns": [
                {"return_type": "Federal", "refund_or_balance": "$452 Refund", "transaction_method": "DD **6095"},
                {"return_type": "State", "refund_or_balance": "Zero Due", "transaction_method": ""},
            ],
            "efiled_returns": [{"name": "Federal"}, {"name": "State"}],
        })
        html = gd.render_template(template, context)
        self.assertEqual(html.count("<tr><td>"), 2)          # one row per return
        self.assertIn("$452 Refund", html)
        self.assertEqual(html.count("<li>"), 2)              # one per e-filed return
        self.assertIn("Our fee for these services is 450.00", html)
        self.assertIn("https://pay.example.com/sam", html)
        self.assertNotIn("{{#", html)                        # no leftover section markers

    def test_results_letter_hides_payment_when_absent(self) -> None:
        template = (gd.REPO_TEMPLATE_DIR / "tax_results_letter.html").read_text(encoding="utf-8")
        html = gd.render_template(template, {"client_name": "Sam", "tax_year": "2025"})
        self.assertNotIn("Our fee", html)                    # no total -> fee line hidden
        self.assertNotIn("pay us securely", html)            # no pay_link -> link hidden

    def test_invoice_total_is_computed(self) -> None:
        context = gd.augment_context(
            {"line_items": [{"amount": "300.00"}, {"amount": "1,200.50"}]}
        )
        self.assertEqual(context["total"], "1,500.50")


class GenerationRunTests(unittest.TestCase):
    def _client(self) -> dict:
        return {
            "firm_name": "Test Firm",
            "client_name": "Jordan Q. Sample",
            "tax_year": "2024",
            "line_items": [{"description": "1040", "amount": "300.00"}],
            "requested_items": [{"item": "W-2 forms"}],
        }

    def test_no_data_file_is_handled(self) -> None:
        with tempfile.TemporaryDirectory() as d:
            result = gd.run_generation(Path(d))
            self.assertEqual(result["document_count"], 0)
            self.assertIn("No clients", result["summary"])

    def test_generates_all_templates_per_client(self) -> None:
        with tempfile.TemporaryDirectory() as d:
            folder = Path(d)
            (folder / "clients.json").write_text(json.dumps([self._client()]), encoding="utf-8")
            result = gd.run_generation(folder)

            self.assertEqual(result["client_count"], 1)
            self.assertEqual(result["document_count"], len(gd.SHIPPED_TEMPLATE_KEYS))
            generated = Path(result["generated_folder"])
            names = {p.name for p in generated.glob("*.html")}
            self.assertIn("Jordan_Q._Sample_invoice.html", names)
            self.assertIn("Jordan_Q._Sample_engagement_letter.html", names)

            invoice = (generated / "Jordan_Q._Sample_invoice.html").read_text()
            self.assertIn("Test Firm", invoice)
            self.assertIn("1040", invoice)          # line item rendered
            self.assertIn("300.00", invoice)        # computed total / amount

    def test_csv_data_file(self) -> None:
        with tempfile.TemporaryDirectory() as d:
            folder = Path(d)
            (folder / "clients.csv").write_text(
                "client_name,tax_year,firm_name\nPat Doe,2024,Test Firm\n", encoding="utf-8"
            )
            result = gd.run_generation(folder, templates=["engagement_letter"])
            self.assertEqual(result["document_count"], 1)
            letter = Path(result["generated_folder"]) / "Pat_Doe_engagement_letter.html"
            self.assertTrue(letter.exists())
            self.assertIn("Pat Doe", letter.read_text())


if __name__ == "__main__":
    unittest.main()
