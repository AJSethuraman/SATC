"""Generate test-fixture documents for the extraction engine.

These are synthetic stand-ins whose KNOWN values let the tests verify that
the extractor pulls the right numbers with correct provenance and dating.
They are fixtures only - the tool's real content comes from documents the
user supplies at run time.

Outputs (in this directory):
  fixture_2013_leveraged_lending.pdf   - interagency guidance excerpt (Type A)
  fixture_2025_rescission.pdf          - OCC/FDIC rescission notice (versioning)
  fixture_2006_cre_guidance.pdf        - CRE concentration guidance excerpt (Type A)
  fixture_sample_cam.docx              - sample credit approval memo (Type B)
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import docx

HERE = Path(__file__).parent

LL_2013 = [
    ("Interagency Guidance on Leveraged Lending", "h"),
    ("Office of the Comptroller of the Currency, Board of Governors of the "
     "Federal Reserve System, and Federal Deposit Insurance Corporation. "
     "Issued March 21, 2013.", "p"),
    ("UNDERWRITING STANDARDS", "h"),
    ("A leverage level after planned asset sales in excess of 6.0x Total Debt to EBITDA "
     "raises concerns for most industries.", "p"),
    ("Institutions should define leveraged lending in their policies. Common definitional "
     "criteria include transactions where the borrower's Total Debt to EBITDA exceeds 4.0x "
     "or Senior Debt to EBITDA exceeds 3.0x.", "p"),
    ("REPAYMENT CAPACITY", "h"),
    ("Base case cash flow projections should show the ability to fully amortize senior secured "
     "debt or repay at least 50 percent of total debt over a five to seven years period.", "p"),
]

RESCISSION_2025 = [
    ("Rescission of the Interagency Guidance on Leveraged Lending", "h"),
    ("December 16, 2025", "p"),
    ("The Office of the Comptroller of the Currency and the Federal Deposit Insurance "
     "Corporation announced today that each agency rescinds the 2013 Interagency Guidance "
     "on Leveraged Lending, effective December 16, 2025.", "p"),
    ("The Board of Governors of the Federal Reserve System did not join this action, and the "
     "guidance remains applicable to institutions supervised by the Federal Reserve.", "p"),
]

CRE_2006 = [
    ("Concentrations in Commercial Real Estate Lending, Sound Risk Management Practices", "h"),
    ("Office of the Comptroller of the Currency, Board of Governors of the Federal Reserve "
     "System, and Federal Deposit Insurance Corporation. Issued December 12, 2006.", "p"),
    ("SUPERVISORY CRITERIA", "h"),
    ("An institution may be identified for further supervisory analysis where total reported loans "
     "for construction, land development, and other land (CLD) represent 100 percent or more of "
     "the institution's total risk-based capital.", "p"),
    ("Further analysis also applies where total non-owner occupied commercial real estate loans "
     "represent 300 percent or more of the institution's total risk-based capital, and the "
     "outstanding balance of the CRE portfolio has increased by 50 percent or more during the "
     "prior 36 months. Loans secured by owner-occupied properties are excluded from this measure.", "p"),
    ("These criteria identify institutions for heightened scrutiny of risk management practices; "
     "they do not constitute limits or hard caps on lending.", "p"),
]

CAM = [
    ("CREDIT APPROVAL MEMORANDUM", "h"),
    ("Borrower: Meridian Fabrication Holdings, LLC", "p"),
    ("Facility: $25,000,000 senior secured revolving credit facility and $40,000,000 term loan A", "p"),
    ("Assigned Risk Grade: 4 (Pass)", "p"),
    ("Guarantor: Mr. Dale Whitfield (founder; personal guaranty, unlimited)", "p"),
    ("Collateral: Blanket lien on all business assets including accounts receivable, inventory, "
     "machinery and equipment", "p"),
    ("FINANCIAL ANALYSIS", "h"),
    ("At close, Total Debt / EBITDA is 3.8x and Senior Debt / EBITDA is 2.9x based on FY2025 "
     "adjusted EBITDA of $17.1 million.", "p"),
    ("DSCR for FY2025 is 1.42x on a consolidated basis; fixed-charge coverage is 1.25x and "
     "interest coverage is 3.6x. The current ratio stands at 1.7x.", "p"),
    ("Global cash flow including guarantor sources produces a global DSCR of 1.55x.", "p"),
    ("REPAYMENT SOURCES", "h"),
    ("The primary source of repayment is operating cash flow from the borrower's fabrication "
     "contracts. The secondary source of repayment is liquidation of pledged working capital "
     "assets under the blanket lien.", "p"),
    ("The facility is personally guaranteed by Mr. Whitfield, whose net worth is reported at "
     "$32 million with liquidity of $4.5 million.", "p"),
    ("COVENANTS", "h"),
    ("Financial covenants include maximum Total Leverage of 4.5x, minimum DSCR of 1.20x, and "
     "minimum tangible net worth of $30 million tested quarterly.", "p"),
    ("PROJECTIONS", "h"),
    ("Management projections assume revenue growth of 8.0% annually and stable gross margins of "
     "31% through FY2028. Projections reflect no new acquisitions during the forecast period.", "p"),
]


def build_docx(spec, path: Path) -> None:
    d = docx.Document()
    style = d.styles["Normal"]
    style.font.name = "Arial"
    for text, kind in spec:
        if kind == "h":
            d.add_heading(text, level=1)
        else:
            d.add_paragraph(text)
    d.save(str(path))


def build_pdf(spec, path: Path) -> None:
    from fpdf import FPDF

    pdf = FPDF(format="letter")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    for text, kind in spec:
        if kind == "h":
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 7, text)
            pdf.ln(2)
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, text)
            pdf.ln(2)
    pdf.output(str(path))


def main() -> None:
    jobs = [
        (LL_2013, "fixture_2013_leveraged_lending", True),
        (RESCISSION_2025, "fixture_2025_rescission", True),
        (CRE_2006, "fixture_2006_cre_guidance", True),
        (CAM, "fixture_sample_cam", False),
    ]
    for spec, stem, as_pdf in jobs:
        if as_pdf:
            build_pdf(spec, HERE / f"{stem}.pdf")
        else:
            build_docx(spec, HERE / f"{stem}.docx")
        print(f"built {stem}{'.pdf' if as_pdf else '.docx'}")


if __name__ == "__main__":
    main()
