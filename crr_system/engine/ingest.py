"""Document ingestion layer.

Normalizes PDF (with OCR fallback hook), Word (.docx), plain/pasted text, and
URLs into a common shape: a list of Page objects, each carrying its text and
1-based page number, so downstream extraction can attach page/section anchors
to every row.
"""

from __future__ import annotations

import io
import re
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional


@dataclass
class Page:
    number: int                 # 1-based
    text: str


@dataclass
class IngestedDocument:
    name: str                   # file name or URL
    kind: str                   # pdf / docx / text / url
    pages: List[Page] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    def full_text(self) -> str:
        return "\n".join(p.text for p in self.pages)


_HEADING_RE = re.compile(
    r"^\s*(?:(?:[IVXLC]+|\d+(?:\.\d+)*)[.)]\s+)?([A-Z][A-Za-z&/\- ]{3,70})\s*$"
)


def detect_section(page_text: str, char_pos: int) -> str:
    """Return the nearest heading-looking line at or above char_pos."""
    head = page_text[:char_pos]
    for line in reversed(head.splitlines()):
        stripped = line.strip()
        if not stripped:
            continue
        if len(stripped) < 72 and (
            stripped.isupper()
            or (stripped.istitle() and not stripped.endswith((".", ",", ";")))
            or _HEADING_RE.match(stripped)
            and stripped.endswith(":")
        ):
            return stripped.rstrip(":")
    return ""


def _ingest_pdf(path: Path) -> IngestedDocument:
    import pdfplumber

    doc = IngestedDocument(name=path.name, kind="pdf")
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                ocr_text, warn = _try_ocr(path, i)
                text = ocr_text
                if warn:
                    doc.warnings.append(warn)
            doc.pages.append(Page(number=i, text=text))
    if not any(p.text.strip() for p in doc.pages):
        doc.warnings.append(
            f"{path.name}: no extractable text on any page (scanned without OCR support?)."
        )
    return doc


def _try_ocr(path: Path, page_number: int) -> tuple:
    """OCR fallback for scanned pages. Returns (text, warning_or_None)."""
    try:
        import pytesseract  # noqa: F401
        from pdf2image import convert_from_path

        images = convert_from_path(
            str(path), first_page=page_number, last_page=page_number, dpi=200
        )
        text = pytesseract.image_to_string(images[0]) if images else ""
        return text, None
    except Exception:
        return "", (
            f"{path.name} p.{page_number}: no text layer and OCR unavailable "
            "(install pytesseract + pdf2image); page skipped."
        )


def _ingest_docx(path: Path) -> IngestedDocument:
    import docx

    d = docx.Document(str(path))
    doc = IngestedDocument(name=path.name, kind="docx")
    # .docx has no fixed pagination pre-render; treat the whole body as one
    # "page" and rely on section headings for location anchors.
    lines = []
    for para in d.paragraphs:
        text = para.text
        if para.style.name.startswith("Heading") and text.strip():
            text = text.strip().upper()  # make headings detectable as sections
        lines.append(text)
    for table in d.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    doc.pages.append(Page(number=1, text="\n".join(lines)))
    return doc


def _ingest_text(source: str, name: str) -> IngestedDocument:
    doc = IngestedDocument(name=name, kind="text")
    # Honor explicit form-feed page breaks in pasted text.
    chunks = source.split("\f") if "\f" in source else [source]
    for i, chunk in enumerate(chunks, start=1):
        doc.pages.append(Page(number=i, text=chunk))
    return doc


def _ingest_url(url: str) -> IngestedDocument:
    req = urllib.request.Request(url, headers={"User-Agent": "crr-extractor/1.0"})
    with urllib.request.urlopen(req, timeout=30) as resp:
        content_type = resp.headers.get("Content-Type", "")
        raw = resp.read()
    if "pdf" in content_type or url.lower().endswith(".pdf"):
        import pdfplumber

        doc = IngestedDocument(name=url, kind="url")
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                doc.pages.append(Page(number=i, text=page.extract_text() or ""))
        return doc
    html = raw.decode("utf-8", errors="replace")
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", "\n", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    doc = IngestedDocument(name=url, kind="url")
    doc.pages.append(Page(number=1, text=text))
    return doc


def ingest(source: str, *, pasted_text: Optional[str] = None) -> IngestedDocument:
    """Ingest a document from a path, URL, or pasted text.

    ingest("policy.pdf")                      -> PDF (OCR fallback if scanned)
    ingest("memo.docx")                       -> Word
    ingest("https://...")                     -> URL (PDF or HTML)
    ingest("pasted", pasted_text="...")       -> raw text, name = first arg
    """
    if pasted_text is not None:
        return _ingest_text(pasted_text, name=source)
    if re.match(r"^https?://", source):
        return _ingest_url(source)
    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(source)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _ingest_pdf(path)
    if suffix in (".docx", ".docm"):
        return _ingest_docx(path)
    if suffix in (".txt", ".md", ""):
        return _ingest_text(path.read_text(encoding="utf-8", errors="replace"), path.name)
    raise ValueError(f"Unsupported document type: {suffix}")
